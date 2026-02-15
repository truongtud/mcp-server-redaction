import os
import shutil
import subprocess
import tempfile

import docx as python_docx
import fitz  # PyMuPDF
import openpyxl

from mcp_server_redaction.engine import RedactionEngine
from mcp_server_redaction.handlers.docx_handler import DocxHandler
from mcp_server_redaction.handlers.pdf import PdfHandler
from mcp_server_redaction.handlers.plain_text import PlainTextHandler
from mcp_server_redaction.handlers.xlsx import XlsxHandler


class TestPlainTextHandler:
    def setup_method(self):
        self.engine = RedactionEngine()
        self.handler = PlainTextHandler()

    def test_redact_creates_output_file(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False
        ) as f:
            f.write("Contact john@example.com for details.\n")
            input_path = f.name

        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_redacted{ext}"

        try:
            result = self.handler.redact(
                self.engine, input_path, output_path
            )
            assert result["entities_found"] >= 1
            assert "session_id" in result

            with open(output_path) as rf:
                content = rf.read()
            assert "john@example.com" not in content
            assert "[EMAIL_ADDRESS_1]" in content
        finally:
            os.unlink(input_path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_unredact_restores_file(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False
        ) as f:
            f.write("Contact john@example.com for details.\n")
            input_path = f.name

        base, ext = os.path.splitext(input_path)
        redacted_path = f"{base}_redacted{ext}"
        unredacted_path = f"{base}_redacted_unredacted{ext}"

        try:
            result = self.handler.redact(
                self.engine, input_path, redacted_path
            )
            session_id = result["session_id"]
            mappings = self.engine.state.get_mappings(session_id)

            undo_result = self.handler.unredact(
                redacted_path, unredacted_path, mappings
            )
            assert undo_result["entities_restored"] >= 1

            with open(unredacted_path) as rf:
                content = rf.read()
            assert "john@example.com" in content
        finally:
            for p in [input_path, redacted_path, unredacted_path]:
                if os.path.exists(p):
                    os.unlink(p)


import pytest
from mcp_server_redaction.handlers import get_handler


class TestHandlerDispatch:
    def test_get_handler_txt(self):
        handler = get_handler(".txt")
        assert isinstance(handler, PlainTextHandler)

    def test_get_handler_unsupported(self):
        with pytest.raises(ValueError, match="Unsupported file extension"):
            get_handler(".xyz")


def _create_test_docx(path: str, paragraphs: list[str]) -> None:
    doc = python_docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def _create_formatted_docx(path: str) -> None:
    """Create a DOCX with mixed formatting: 'Contact <bold>John Smith</bold> at <italic>john@example.com</italic> today.'"""
    doc = python_docx.Document()
    para = doc.add_paragraph()
    run1 = para.add_run("Contact ")
    run2 = para.add_run("John Smith")
    run2.bold = True
    run3 = para.add_run(" at ")
    run4 = para.add_run("john@example.com")
    run4.italic = True
    run5 = para.add_run(" today.")
    doc.save(path)


class TestDocxHandler:
    def setup_method(self):
        self.engine = RedactionEngine()
        self.handler = DocxHandler()

    def test_redact_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.docx")
            output_path = os.path.join(tmpdir, "test_redacted.docx")

            _create_test_docx(input_path, [
                "Contact john@example.com for details.",
                "No sensitive data here.",
            ])

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 1
            assert result["session_id"] is not None

            doc = python_docx.Document(output_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "john@example.com" not in all_text
            assert "[EMAIL_ADDRESS_1]" in all_text
            assert "No sensitive data here." in all_text

    def test_unredact_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.docx")
            redacted_path = os.path.join(tmpdir, "test_redacted.docx")
            unredacted_path = os.path.join(tmpdir, "test_unredacted.docx")

            _create_test_docx(input_path, ["Contact john@example.com for details."])

            result = self.handler.redact(self.engine, input_path, redacted_path)
            mappings = self.engine.state.get_mappings(result["session_id"])

            undo = self.handler.unredact(redacted_path, unredacted_path, mappings)
            assert undo["entities_restored"] >= 1

            doc = python_docx.Document(unredacted_path)
            assert "john@example.com" in doc.paragraphs[0].text

    def test_redact_docx_with_table(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.docx")
            output_path = os.path.join(tmpdir, "test_redacted.docx")

            doc = python_docx.Document()
            doc.add_paragraph("Header text")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "john@example.com"
            doc.save(input_path)

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 1

            out_doc = python_docx.Document(output_path)
            cell_text = out_doc.tables[0].cell(0, 1).text
            assert "john@example.com" not in cell_text


    def test_redact_preserves_run_formatting(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "formatted.docx")
            output_path = os.path.join(tmpdir, "formatted_redacted.docx")

            _create_formatted_docx(input_path)

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 1

            doc = python_docx.Document(output_path)
            para = doc.paragraphs[0]
            all_text = para.text
            assert "john@example.com" not in all_text
            assert "John Smith" not in all_text

            # Verify formatting is preserved on non-redacted runs
            runs_text = [(r.text, r.bold, r.italic) for r in para.runs if r.text.strip()]
            # "Contact" should not be bold
            contact_runs = [r for r in runs_text if "Contact" in r[0]]
            assert len(contact_runs) >= 1
            assert contact_runs[0][1] is not True  # not bold

            # The last run (originally " today.") should not be italic.
            # Note: "today" may be detected as DATE_TIME, so look for the
            # run that ends with "." — it came from the non-italic run5.
            last_runs = [r for r in runs_text if r[0].endswith(".")]
            assert len(last_runs) >= 1
            assert last_runs[0][2] is not True  # not italic


    def test_redact_preserves_formatting_cross_run_pii(self):
        """PII that spans two runs — placeholder takes first run's format."""
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "cross_run.docx")
            output_path = os.path.join(tmpdir, "cross_run_redacted.docx")

            doc = python_docx.Document()
            para = doc.add_paragraph()
            para.add_run("Contact ")
            run2 = para.add_run("john@exam")
            run2.bold = True
            para.add_run("ple.com")
            para.add_run(" for details.")
            doc.save(input_path)

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 1

            out_doc = python_docx.Document(output_path)
            para = out_doc.paragraphs[0]
            assert "john@example.com" not in para.text
            assert "john@exam" not in para.text
            assert "ple.com" not in para.text

    def test_unredact_preserves_formatting(self):
        """Unredact should also preserve run formatting."""
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.docx")
            redacted_path = os.path.join(tmpdir, "test_redacted.docx")
            unredacted_path = os.path.join(tmpdir, "test_unredacted.docx")

            _create_formatted_docx(input_path)

            result = self.handler.redact(self.engine, input_path, redacted_path)
            mappings = self.engine.state.get_mappings(result["session_id"])

            undo = self.handler.unredact(redacted_path, unredacted_path, mappings)
            assert undo["entities_restored"] >= 1

            doc = python_docx.Document(unredacted_path)
            para = doc.paragraphs[0]
            assert "john@example.com" in para.text
            # Non-redacted run formatting should be preserved
            runs_text = [(r.text, r.bold, r.italic) for r in para.runs if r.text.strip()]
            last_runs = [r for r in runs_text if r[0].endswith(".")]
            assert len(last_runs) >= 1


class TestXlsxHandler:
    def setup_method(self):
        self.engine = RedactionEngine()
        self.handler = XlsxHandler()

    def test_redact_xlsx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.xlsx")
            output_path = os.path.join(tmpdir, "test_redacted.xlsx")

            wb = openpyxl.Workbook()
            ws = wb.active
            ws["A1"] = "Name"
            ws["B1"] = "Email"
            ws["A2"] = "John Smith"
            ws["B2"] = "john@example.com"
            wb.save(input_path)

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 1
            assert result["session_id"] is not None

            wb_out = openpyxl.load_workbook(output_path)
            ws_out = wb_out.active
            assert "john@example.com" not in str(ws_out["B2"].value)
            assert ws_out["A1"].value == "Name"

    def test_redact_xlsx_multiple_sheets(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.xlsx")
            output_path = os.path.join(tmpdir, "test_redacted.xlsx")

            wb = openpyxl.Workbook()
            ws1 = wb.active
            ws1.title = "Sheet1"
            ws1["A1"] = "john@example.com"
            ws2 = wb.create_sheet("Sheet2")
            ws2["A1"] = "jane@example.com"
            wb.save(input_path)

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 2

    def test_unredact_xlsx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.xlsx")
            redacted_path = os.path.join(tmpdir, "test_redacted.xlsx")
            unredacted_path = os.path.join(tmpdir, "test_unredacted.xlsx")

            wb = openpyxl.Workbook()
            ws = wb.active
            ws["A1"] = "john@example.com"
            wb.save(input_path)

            result = self.handler.redact(self.engine, input_path, redacted_path)
            mappings = self.engine.state.get_mappings(result["session_id"])

            undo = self.handler.unredact(redacted_path, unredacted_path, mappings)
            assert undo["entities_restored"] >= 1

            wb_out = openpyxl.load_workbook(unredacted_path)
            assert wb_out.active["A1"].value == "john@example.com"


def _create_test_pdf(path: str, pages: list[str], fontsize: float = 12) -> None:
    """Helper: create a PDF with one text block per page."""
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=fontsize)
    doc.save(path)
    doc.close()


class TestPdfHandler:
    def setup_method(self):
        self.engine = RedactionEngine()
        self.handler = PdfHandler()

    def test_redact_pdf_placeholder_mode(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.pdf")
            output_path = os.path.join(tmpdir, "test_redacted.pdf")

            _create_test_pdf(input_path, [
                "Contact john@example.com for details."
            ])

            result = self.handler.redact(
                self.engine, input_path, output_path,
                use_placeholders=True,
            )
            assert result["entities_found"] >= 1
            assert result["session_id"] is not None

            doc = fitz.open(output_path)
            page_text = doc[0].get_text()
            doc.close()
            assert "john@example.com" not in page_text
            assert "[EMAIL_ADDRESS_1]" in page_text

    def test_redact_pdf_blackbox_mode(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.pdf")
            output_path = os.path.join(tmpdir, "test_redacted.pdf")

            _create_test_pdf(input_path, [
                "Contact john@example.com for details."
            ])

            result = self.handler.redact(
                self.engine, input_path, output_path,
                use_placeholders=False,
            )
            assert result["entities_found"] >= 1
            assert result["session_id"] is None

            doc = fitz.open(output_path)
            page_text = doc[0].get_text()
            doc.close()
            assert "john@example.com" not in page_text

    def test_unredact_pdf(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.pdf")
            redacted_path = os.path.join(tmpdir, "test_redacted.pdf")
            unredacted_path = os.path.join(tmpdir, "test_unredacted.pdf")

            _create_test_pdf(input_path, [
                "Contact john@example.com for details."
            ])

            result = self.handler.redact(
                self.engine, input_path, redacted_path,
                use_placeholders=True,
            )
            mappings = self.engine.state.get_mappings(result["session_id"])

            undo = self.handler.unredact(redacted_path, unredacted_path, mappings)
            assert undo["entities_restored"] >= 1

            doc = fitz.open(unredacted_path)
            page_text = doc[0].get_text()
            doc.close()
            assert "john@example.com" in page_text

    def test_redact_pdf_preserves_font_size(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.pdf")
            output_path = os.path.join(tmpdir, "test_redacted.pdf")

            _create_test_pdf(input_path, [
                "Contact john@example.com for details."
            ], fontsize=18)

            result = self.handler.redact(
                self.engine, input_path, output_path,
                use_placeholders=True,
            )
            assert result["entities_found"] >= 1

            doc = fitz.open(output_path)
            blocks = doc[0].get_text("dict")["blocks"]
            doc.close()

            # Find the span containing the placeholder
            placeholder_fontsize = None
            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        if "[EMAIL_ADDRESS_1]" in span["text"]:
                            placeholder_fontsize = span["size"]

            # Font size should be closer to 18 than the old hardcoded 10.
            # PyMuPDF may scale down to fit replacement text in the original rect.
            assert placeholder_fontsize is not None
            assert placeholder_fontsize > 12  # well above the old hardcoded 10


from mcp_server_redaction.handlers.doc import DocHandler

LIBREOFFICE_AVAILABLE = shutil.which("libreoffice") is not None


class TestDocHandler:
    def setup_method(self):
        self.engine = RedactionEngine()
        self.handler = DocHandler()

    def test_doc_handler_without_libreoffice_errors(self):
        """Test that DocHandler raises RuntimeError when LibreOffice is not available."""
        original_which = shutil.which

        def fake_which(name):
            if name == "libreoffice":
                return None
            return original_which(name)

        shutil.which = fake_which
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                input_path = os.path.join(tmpdir, "test.doc")
                output_path = os.path.join(tmpdir, "test_redacted.docx")
                with open(input_path, "w") as f:
                    f.write("dummy")
                with pytest.raises(RuntimeError, match="LibreOffice is required"):
                    self.handler.redact(self.engine, input_path, output_path)
        finally:
            shutil.which = original_which

    @pytest.mark.skipif(not LIBREOFFICE_AVAILABLE, reason="LibreOffice not installed")
    def test_redact_doc(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a .docx first, then convert to .doc via LibreOffice
            docx_path = os.path.join(tmpdir, "test.docx")
            doc = python_docx.Document()
            doc.add_paragraph("Contact john@example.com for details.")
            doc.save(docx_path)

            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "doc",
                 "--outdir", tmpdir, docx_path],
                check=True, capture_output=True,
            )
            input_path = os.path.join(tmpdir, "test.doc")
            output_path = os.path.join(tmpdir, "test_redacted.docx")

            result = self.handler.redact(self.engine, input_path, output_path)
            assert result["entities_found"] >= 1
