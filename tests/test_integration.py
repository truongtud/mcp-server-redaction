import json
import os
import tempfile

import docx as python_docx
import fitz
import openpyxl

from mcp_server_redaction.engine import RedactionEngine
from mcp_server_redaction.tools import (
    handle_redact,
    handle_unredact,
    handle_analyze,
    handle_configure,
)
from mcp_server_redaction.tools.redact_file import handle_redact_file
from mcp_server_redaction.tools.unredact_file import handle_unredact_file


class TestEndToEnd:
    def setup_method(self):
        self.engine = RedactionEngine()

    def test_full_redact_unredact_cycle(self):
        original = (
            "John Smith's email is john@example.com "
            "and you can call him at 555-123-4567."
        )

        # Redact
        redact_result = json.loads(handle_redact(self.engine, text=original))
        redacted = redact_result["redacted_text"]
        session_id = redact_result["session_id"]

        assert "john@example.com" not in redacted
        assert "John Smith" not in redacted
        assert redact_result["entities_found"] >= 2

        # Unredact
        unredact_result = json.loads(
            handle_unredact(self.engine, redacted_text=redacted, session_id=session_id)
        )
        assert unredact_result["original_text"] == original

    def test_analyze_then_selective_redact(self):
        text = "Dr. Jane Doe prescribed Metformin for patient MRN: 123-456-789"

        # First analyze to see what's in the text
        analysis = json.loads(handle_analyze(self.engine, text=text))
        found_types = {e["type"] for e in analysis["entities"]}
        assert len(found_types) >= 1

        # Then redact only PERSON entities
        redact_result = json.loads(
            handle_redact(self.engine, text=text, entity_types=["PERSON"])
        )
        # PERSON should be redacted
        assert "[PERSON_" in redact_result["redacted_text"]

    def test_configure_custom_pattern_then_redact(self):
        # Add a custom pattern
        config_result = json.loads(
            handle_configure(
                self.engine,
                custom_patterns=[
                    {"name": "PROJECT_CODE", "pattern": r"PRJ-\d{4}", "score": 0.95}
                ],
            )
        )
        assert "PROJECT_CODE" in config_result["active_entities"]

        # Now redact text containing the custom pattern
        text = "Assign this to PRJ-1234 immediately"
        redact_result = json.loads(handle_redact(self.engine, text=text))
        assert "[PROJECT_CODE_1]" in redact_result["redacted_text"]
        assert "PRJ-1234" not in redact_result["redacted_text"]


class TestHybridDetectionIntegration:
    def setup_method(self):
        self.engine = RedactionEngine(use_llm=False)

    def test_redact_english_pii_comprehensive(self):
        text = (
            "Patient John Smith (DOB: 03/15/1985) visited Dr. Sarah Johnson. "
            "Insurance: POL-2024-00045678. Email: john.smith@hospital.org. "
            "Prescribed Metformin 500mg. NPI: 1234567890."
        )
        result = self.engine.redact(text)
        assert result["entities_found"] >= 4  # name, email, drug, policy at minimum
        assert "john.smith@hospital.org" not in result["redacted_text"]
        assert "John Smith" not in result["redacted_text"]

    def test_redact_german_text(self):
        text = "Herr Hans Müller wohnt in der Berliner Straße 42, 10115 Berlin. Tel: +49 30 12345678."
        result = self.engine.redact(text)
        assert result["entities_found"] >= 1  # GLiNER should catch the name at minimum

    def test_redact_mixed_language(self):
        text = (
            "Customer Nguyễn Văn An called about policy POL-2024-00099999. "
            "His email is an.nguyen@example.com."
        )
        result = self.engine.redact(text)
        assert "an.nguyen@example.com" not in result["redacted_text"]
        assert result["entities_found"] >= 2

    def test_unredact_still_works_after_hybrid_detection(self):
        text = "Contact Jane Doe at jane@example.com"
        redact_result = self.engine.redact(text)
        unredact_result = self.engine.unredact(
            redact_result["redacted_text"],
            redact_result["session_id"],
        )
        assert "jane@example.com" in unredact_result["original_text"]

    def test_file_redaction_uses_hybrid_engine(self, tmp_path):
        """File handlers should benefit from the hybrid engine automatically."""
        test_file = tmp_path / "test.txt"
        test_file.write_text("Send invoice to Hans Müller, hans@firma.de, policy POL-2024-00012345")

        result = json.loads(
            handle_redact_file(self.engine, str(test_file))
        )
        assert result["entities_found"] >= 2
        assert "error" not in result


class TestFileFormatRoundtrips:
    def setup_method(self):
        self.engine = RedactionEngine()

    def test_docx_roundtrip(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.docx")
            doc = python_docx.Document()
            doc.add_paragraph("Contact john@example.com for details.")
            doc.save(input_path)

            redact_result = json.loads(
                handle_redact_file(self.engine, file_path=input_path)
            )
            assert redact_result["entities_found"] >= 1

            unredact_result = json.loads(
                handle_unredact_file(
                    self.engine,
                    file_path=redact_result["redacted_file_path"],
                    session_id=redact_result["session_id"],
                )
            )
            assert unredact_result["entities_restored"] >= 1

            doc_out = python_docx.Document(unredact_result["unredacted_file_path"])
            assert "john@example.com" in doc_out.paragraphs[0].text

    def test_xlsx_roundtrip(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.xlsx")
            wb = openpyxl.Workbook()
            wb.active["A1"] = "john@example.com"
            wb.save(input_path)

            redact_result = json.loads(
                handle_redact_file(self.engine, file_path=input_path)
            )
            assert redact_result["entities_found"] >= 1

            unredact_result = json.loads(
                handle_unredact_file(
                    self.engine,
                    file_path=redact_result["redacted_file_path"],
                    session_id=redact_result["session_id"],
                )
            )
            assert unredact_result["entities_restored"] >= 1

            wb_out = openpyxl.load_workbook(unredact_result["unredacted_file_path"])
            assert wb_out.active["A1"].value == "john@example.com"

    def test_pdf_placeholder_roundtrip(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "test.pdf")
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Contact john@example.com for details.", fontsize=12)
            doc.save(input_path)
            doc.close()

            redact_result = json.loads(
                handle_redact_file(
                    self.engine, file_path=input_path, use_placeholders=True
                )
            )
            assert redact_result["entities_found"] >= 1

            unredact_result = json.loads(
                handle_unredact_file(
                    self.engine,
                    file_path=redact_result["redacted_file_path"],
                    session_id=redact_result["session_id"],
                )
            )
            assert unredact_result["entities_restored"] >= 1

            doc_out = fitz.open(unredact_result["unredacted_file_path"])
            assert "john@example.com" in doc_out[0].get_text()
            doc_out.close()


class TestConfigureThreshold:
    def setup_method(self):
        self.engine = RedactionEngine()

    def test_configure_sets_score_threshold(self):
        result = json.loads(
            handle_configure(self.engine, score_threshold=0.7)
        )
        assert result["status"] == "ok"
        assert result["score_threshold"] == 0.7
        assert self.engine.score_threshold == 0.7

    def test_configure_reports_current_threshold(self):
        result = json.loads(handle_configure(self.engine))
        assert "score_threshold" in result
        assert result["score_threshold"] == 0.4  # default


class TestFalsePositiveRegression:
    """Regression tests for known false positive patterns.

    These sentences previously triggered false detections. They should
    produce zero or minimal redactions with the precision improvements.
    """

    def setup_method(self):
        self.engine = RedactionEngine(use_llm=False)

    def test_blog_prose_not_over_redacted(self):
        text = (
            "I Built an MCP Server That Lets Claude Redact Your Documents "
            "Before Reading Them. The gap between 'I need an LLM' and "
            "'this data shouldn't leave my machine' is wider than it should be."
        )
        result = self.engine.redact(text)
        # This prose contains no real PII — should have very few or zero detections.
        # "Claude" might still be flagged as PERSON (acceptable), but nothing else.
        assert result["entities_found"] <= 2
        assert "SWIFT_CODE" not in result["redacted_text"]

    def test_common_words_not_flagged_as_swift(self):
        text = "The credentials in the document are separate from the database."
        result = self.engine.redact(text)
        assert "SWIFT_CODE" not in result["redacted_text"]

    def test_time_expressions_not_over_flagged(self):
        text = "They spend twenty minutes manually scrubbing names and account numbers."
        result = self.engine.redact(text)
        # "twenty minutes" should not become [DATE_TIME_N]
        assert "twenty minutes" in result["redacted_text"] or result["entities_found"] <= 1

    def test_product_names_not_flagged(self):
        text = "Something where Claude could clean the document itself."
        result = self.engine.redact(text)
        assert "SWIFT_CODE" not in result["redacted_text"]

    def test_real_pii_still_detected(self):
        """Ensure precision improvements don't kill recall on real PII."""
        text = (
            "Patient John Smith (DOB: 03/15/1985) visited Dr. Sarah Johnson. "
            "Email: john.smith@hospital.org. Insurance: POL-2024-00045678."
        )
        result = self.engine.redact(text)
        assert "john.smith@hospital.org" not in result["redacted_text"]
        assert "John Smith" not in result["redacted_text"]
        assert result["entities_found"] >= 3
