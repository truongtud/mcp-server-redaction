import docx as python_docx

from ..engine import RedactionEngine
from .base import FileHandler


class DocxHandler(FileHandler):
    def redact(
        self,
        engine: RedactionEngine,
        input_path: str,
        output_path: str,
        entity_types: list[str] | None = None,
        use_placeholders: bool = True,
    ) -> dict:
        doc = python_docx.Document(input_path)
        total_found = 0
        session_id = None

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            result = engine.redact(para.text, entity_types=entity_types)
            if result["entities_found"] > 0:
                total_found += result["entities_found"]
                if session_id is None:
                    session_id = result["session_id"]
                else:
                    self._merge_session(engine, session_id, result["session_id"])
                self._surgical_replace(para, result["entities"])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        result = engine.redact(para.text, entity_types=entity_types)
                        if result["entities_found"] > 0:
                            total_found += result["entities_found"]
                            if session_id is None:
                                session_id = result["session_id"]
                            else:
                                self._merge_session(engine, session_id, result["session_id"])
                            self._surgical_replace(para, result["entities"])

        if session_id is None:
            session_id = engine.state.create_session()

        doc.save(output_path)
        return {"session_id": session_id, "entities_found": total_found}

    def unredact(
        self,
        input_path: str,
        output_path: str,
        mappings: dict[str, str],
    ) -> dict:
        doc = python_docx.Document(input_path)
        entities_restored = 0

        for para in doc.paragraphs:
            count = self._surgical_unredact(para, mappings)
            entities_restored += count

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count = self._surgical_unredact(para, mappings)
                        entities_restored += count

        doc.save(output_path)
        return {"entities_restored": entities_restored}

    @staticmethod
    def _surgical_replace(para, entities: list[dict]) -> None:
        """Replace PII in paragraph runs surgically, preserving formatting.

        Falls back to full-paragraph replacement if run mapping fails.
        """
        if not entities:
            return

        runs = para.runs
        if not runs:
            # No runs — simple text-only paragraph, just replace
            text = para.text
            for ent in sorted(entities, key=lambda e: e["original_start"], reverse=True):
                text = text[:ent["original_start"]] + ent["placeholder"] + text[ent["original_end"]:]
            para.text = text
            return

        # Build run offset map
        concatenated = "".join(r.text for r in runs)
        if concatenated != para.text:
            # Run text doesn't match paragraph text — fall back
            text = para.text
            for ent in sorted(entities, key=lambda e: e["original_start"], reverse=True):
                text = text[:ent["original_start"]] + ent["placeholder"] + text[ent["original_end"]:]
            runs[0].text = text
            for r in runs[1:]:
                r.text = ""
            return

        # Map each run to its character range
        run_ranges = []
        offset = 0
        for run in runs:
            end = offset + len(run.text)
            run_ranges.append((offset, end))
            offset = end

        # Process entities right-to-left to preserve positions
        for ent in sorted(entities, key=lambda e: e["original_start"], reverse=True):
            orig_start = ent["original_start"]
            orig_end = ent["original_end"]
            placeholder = ent["placeholder"]

            for i, (run_start, run_end) in enumerate(run_ranges):
                if orig_start >= run_end:
                    continue
                if orig_start < run_start:
                    break

                local_start = orig_start - run_start

                if orig_end <= run_end:
                    # Case 1: PII fits entirely within this single run
                    local_end = orig_end - run_start
                    runs[i].text = runs[i].text[:local_start] + placeholder + runs[i].text[local_end:]
                else:
                    # Case 2: PII crosses run boundaries
                    runs[i].text = runs[i].text[:local_start] + placeholder
                    for j in range(i + 1, len(run_ranges)):
                        sub_start, sub_end = run_ranges[j]
                        if sub_end <= orig_end:
                            runs[j].text = ""
                        else:
                            local_trim = orig_end - sub_start
                            runs[j].text = runs[j].text[local_trim:]
                            break
                break

    @staticmethod
    def _surgical_unredact(para, mappings: dict[str, str]) -> int:
        """Replace placeholders in runs surgically, preserving formatting."""
        count = 0
        for run in para.runs:
            for placeholder, original in mappings.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, original)
                    count += 1
        if count == 0:
            # Check if placeholder spans runs (unlikely but possible)
            full_text = para.text
            for placeholder, original in mappings.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, original)
                    count += 1
            if count > 0 and para.runs:
                para.runs[0].text = full_text
                for r in para.runs[1:]:
                    r.text = ""
        return count

    @staticmethod
    def _apply_mappings(text: str, mappings: dict[str, str]) -> tuple[str, int]:
        count = 0
        for placeholder, original in mappings.items():
            if placeholder in text:
                text = text.replace(placeholder, original)
                count += 1
        return text, count

    @staticmethod
    def _merge_session(engine: RedactionEngine, target_id: str, source_id: str) -> None:
        source_mappings = engine.state.get_mappings(source_id)
        if source_mappings:
            for placeholder, original in source_mappings.items():
                engine.state.add_mapping(target_id, placeholder, original)
